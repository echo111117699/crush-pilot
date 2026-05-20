import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Title
title = doc.add_heading('Crush Pilot 产品草图作业', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Meta info
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run('提交人：Echo Li　　日期：2026-05-19　　版本：原型 v1.0')
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Helper functions
def add_h1(text):
    h = doc.add_heading(text, level=1)
    return h

def add_h2(text):
    h = doc.add_heading(text, level=2)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_bold_body(bold_part, normal_part):
    p = doc.add_paragraph()
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(normal_part)
    return p

def add_image(path, width_inches=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_inches))

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===== Document Content =====

add_h1('一、产品概述')
add_body('Crush Pilot 是一款帮助年轻女性在暧昧关系中保持清醒、获得主体性的分析工具。通过上传聊天截图，AI 解读双方的行为模式，积累动态人格洞察，并在关键决策点提供高价值策略。')
add_body('一句话定义：通过持续积累的对话数据，帮用户看清一个人，也看清自己。')
add_body('目标用户：18-28 岁女性，正处于暧昧/不确定关系中，希望用理性数据而非情绪冲动来做关系决策。')

add_h1('二、产品价值')

add_h2('2.1 为什么做这个产品')
add_body('市面上的聊天分析工具有两个核心缺陷：')

add_table(
    ['现有产品的问题', 'Crush Pilot 的解法'],
    [
        ['没有人格档案的积累机制——每次分析都是"一次性"的，无法形成对一个人的系统性认知',
         '"他"线：从每次分析中提取行为证据，累积成动态人格洞察。第一次发现"闭合式回应"，第二次印证"回避性信号"——认知随数据增长而清晰'],
        ['过度提供情绪价值，帮用户"自我安慰"，缺乏清醒度和主体性',
         '"我"线：用数据帮用户意识到自己在关系里的模式（主动程度、情绪波动），保持清醒。不含星座、塔罗牌等泛娱乐化内容'],
    ]
)

add_h2('2.2 两条流动的线')
add_body('产品中的所有功能都围绕两条核心数据线组织：')
add_bold_body('"他"线（关于他）：', '上传截图 → AI 分析 → 提取行为模式 → 积累人格洞察 → 策略建议')
add_bold_body('"我"线（关于你）：', '自我档案 → 使用数据 → 关系模式发现 → 静默期工具 → 每周高光')
add_body('这两条线不是割裂的，而是通过同一次分析同时更新——你在看清他的同时，也在看清自己。')

add_h1('三、设计思路')

add_h2('3.1 信息架构：从 3 Tab 到 2 Tab')
add_body('原始设计按"功能类型"分为聊天/档案/我三个 Tab，功能完整但割裂。重构后按"用户任务流"分为 2 个 Tab：')

add_table(
    ['Tab', '用户任务', '对应数据线'],
    [
        ['分析', '上传聊天 · 获取策略', '"他"线的入口'],
        ['洞察', '人格档案 · 关系模式', '"他"线 + "我"线的沉淀'],
    ]
)

add_h2('3.2 冷启动精简')
add_body('旧流程：Welcome → Questionnaire（10题）→ Transition → CrushForm → Main，4 步才进入主界面，问卷填完即封存，感受不到与后续分析的连接。')
add_body('新流程：Welcome → QuickContext（6题一屏）→ Main，2 步进入主界面。')

add_table(
    ['问题', '为什么问'],
    [
        ['你的称呼？他的称呼？', 'AI 怎么称呼你和指代他'],
        ['怎么认识的？多久了？', '关系背景，影响策略风格'],
        ['你现在最想知道什么？', '直接影响分析方向'],
        ['你在这段关系里最常有的情绪？', '帮你认识自己的模式'],
        ['还有什么要告诉我的？', '开放入口，补充上下文'],
    ]
)

add_h2('3.3 数据流动设计')
add_body('分析不再是一次性的——每次上传截图后，系统从分析结果中提取结构化洞察，存入全局状态，在"洞察"页持续积累：')
add_body('上传截图 → AI 分析 → { 护盾语 + 客观看待（事实还原/情绪解读/清醒提醒）+ 高价值策略 + 行为洞察 } → 洞察页累积展示')

add_h2('3.4 视觉系统')
add_body('配色：马卡龙色系（粉 #EEACC2 / 黄 #F3E49B / 蓝 #9DD9E1），柔和但不过度甜美')
add_body('排版：行距全部 ≥ 1.9，文字密度低，留白充足')
add_body('卡片设计：圆角 14px+，轻阴影，统一边框色，信息层级通过背景色区分')
add_body('对话框：固定宽度 400px，保证长文本换行时各气泡宽度一致')

add_h2('3.5 关键交互设计')
add_body('静默期时钟：启动静默后，页面显示 SVG 圆形倒计时，精确到小时和分钟，提供 6 种替代活动建议')
add_body('策略滑块：三段式滑动选择（试探 ↔ 适中 ↔ 直接），含淡入淡出动画和一键复制功能')
add_body('洞察 Tab：统一白色卡片包裹，内部用分割线区分功能区，用 Tab 切换子视图')

add_h1('四、关键界面')

screenshot_dir = '/Users/echoli/crush-pilot/screenshots'

add_h2('4.1 隐私声明')
add_body('用户首次打开 App 时看到。强调数据安全——云端分析后即销毁、档案仅本地存储、不向第三方透露。')
add_image(f'{screenshot_dir}/01-privacy.png')

add_h2('4.2 欢迎页')
add_body('阐明产品定位：不看星盘、不算塔罗、基于真实对话数据的理性分析。')
add_image(f'{screenshot_dir}/02-welcome.png')

add_h2('4.3 冷启动问卷（QuickContext）')
add_body('6 个问题一屏完成，代替传统的 10 题问卷 + CrushForm。填写内容直接影响后续 AI 分析的方向和语气。')
add_image(f'{screenshot_dir}/03-quickcontext.png')

add_h2('4.4 分析页')
add_body('顶部护盾卡片 + 对话流 + 底部上传区。上传截图后 AI 返回：护盾 & 客观分析 + 高价值策略（含滑块切换）。')
add_image(f'{screenshot_dir}/04-analyze.png')

add_h2('4.5 洞察页')
add_body('两个核心板块——"关于他"：可编辑档案 + 行为模式发现；"关于你"：三个子 Tab（你的状态 / 静默期 / 这周的你）。')
add_image(f'{screenshot_dir}/05-insight.png')

add_h1('五、技术实现')

add_table(
    ['层', '技术选型'],
    [
        ['前端框架', 'React 19 + Vite 6'],
        ['样式', 'Tailwind CSS v4，自定义主题色'],
        ['状态管理', 'React Context API（AppContext）'],
        ['AI 接入', 'DeepSeek Chat API（deepseek-chat 模型）'],
        ['数据持久化', 'localStorage（用户档案、洞察积累、静默历史）'],
        ['OCR（规划中）', 'Tesseract.js 浏览器端识别'],
    ]
)

add_h1('六、后续规划')

add_body('1. OCR 真实截图识别：接入 Tesseract.js 或云端 OCR，让用户上传真实聊天截图')
add_body('2. 更长周期的洞察积累：基于 10+ 次分析后的趋势报告')
add_body('3. 更多行为模式维度：从回应节奏、语言模式，扩展到情感温度、话题主导权等')
add_body('4. 导出报告：生成可分享/保存的 PDF 关系洞察报告')

# Save
output_path = '/Users/echoli/crush-pilot/docs/Crush_Pilot_草图作业.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
